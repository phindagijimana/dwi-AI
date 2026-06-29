"""Convert a markdown file to a .docx via python-docx.

Focused converter — handles only the constructs used in ``nodestrength.md``:

* ATX headings (``#`` .. ``####``)
* Paragraphs with inline ``code``, **bold**, *italic*, and ``[text](url)`` links
* Bullet lists (``-`` or ``*``)
* Numbered lists (``1.``)
* Tables (pipe-delimited with a ``|---|---|`` separator row)
* Fenced code blocks (``` ``` ``` ``)
* Horizontal rules (``---``)

No frontmatter parsing, no nested lists, no images. Adequate for technical
documentation.

Usage:
    python scripts/md_to_docx.py --in nodestrength.md --out nodestrength.docx
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, RGBColor


_INLINE_RE = re.compile(
    r"(`[^`]+`|"            # `code`
    r"\*\*[^*]+\*\*|"        # **bold**
    r"\*[^*\n]+\*|"          # *italic*
    r"\[[^\]]+\]\([^)]+\))"  # [text](url)
)


def _add_inline_runs(paragraph, text: str) -> None:
    """Add `text` to `paragraph` as runs, respecting inline markdown."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Plain text before this match.
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                link_text = link_match.group(1)
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
                run.font.underline = True
            else:
                paragraph.add_run(token)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_separator(line: str) -> bool:
    s = line.strip()
    if not s.startswith("|") or not s.endswith("|"):
        return False
    inner = s.strip("|")
    cells = [c.strip() for c in inner.split("|")]
    return all(re.fullmatch(r":?-+:?", c) for c in cells if c)


def _split_table_row(line: str) -> List[str]:
    s = line.strip().strip("|")
    return [c.strip() for c in s.split("|")]


def convert(md_text: str, doc: Document) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.lstrip().startswith("```"):
            i += 1
            buf: List[str] = []
            while i < len(lines) and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run("\n".join(buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        # Heading
        m = re.match(r"^(#{1,4})\s+(.*?)\s*$", line)
        if m:
            level = len(m.group(1))
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.add_run("─" * 50)
            i += 1
            continue

        # Table
        if line.strip().startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header_cells = _split_table_row(line)
            i += 2  # skip separator
            body: List[List[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                body.append(_split_table_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1 + len(body), cols=len(header_cells))
            tbl.style = "Light Grid Accent 1"
            for j, cell_text in enumerate(header_cells):
                cell = tbl.rows[0].cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.bold = True
            for r, row_cells in enumerate(body, start=1):
                for j, cell_text in enumerate(row_cells):
                    if j >= len(tbl.rows[r].cells):
                        continue
                    cell = tbl.rows[r].cells[j]
                    p = cell.paragraphs[0]
                    _add_inline_runs(p, cell_text)
            doc.add_paragraph()
            continue

        # Bullet list
        if re.match(r"^\s*[-*]\s+", line):
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                item = re.sub(r"^\s*[-*]\s+", "", lines[i])
                p = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(p, item)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            while i < len(lines) and re.match(r"^\s*\d+\.\s+", lines[i]):
                item = re.sub(r"^\s*\d+\.\s+", "", lines[i])
                p = doc.add_paragraph(style="List Number")
                _add_inline_runs(p, item)
                i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Regular paragraph — collect contiguous non-special lines.
        para_lines: List[str] = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            if (nxt.strip() == ""
                or nxt.lstrip().startswith("```")
                or re.match(r"^#{1,4}\s", nxt)
                or re.match(r"^\s*[-*]\s+", nxt)
                or re.match(r"^\s*\d+\.\s+", nxt)
                or nxt.strip().startswith("|")
                or nxt.strip() == "---"):
                break
            para_lines.append(nxt)
            i += 1
        text = " ".join(s.strip() for s in para_lines)
        p = doc.add_paragraph()
        _add_inline_runs(p, text)


def main(argv=None) -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--in", dest="src", required=True, type=Path)
    ap.add_argument("--out", required=True, type=Path)
    args = ap.parse_args(argv)

    md = args.src.read_text(encoding="utf-8")
    doc = Document()

    # Default body style: a small bump in size for readability.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    convert(md, doc)
    args.out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.out))
    print(f"Wrote {args.out}  ({args.out.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
